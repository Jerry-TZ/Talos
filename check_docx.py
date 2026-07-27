from docx import Document
doc = Document('report.docx')
print('表格行数:', len(doc.tables[0].rows))
for i, row in enumerate(doc.tables[0].rows):
    print(f'行{i}:')
    for j, cell in enumerate(row.cells):
        print(f'  列{j}: [{cell.text[:60]}]')
